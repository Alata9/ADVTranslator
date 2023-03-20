import os
from tkinter import *
from tkinter import Checkbutton, messagebox, scrolledtext, filedialog, ttk
import pyttsx3
from googletrans import Translator
import docx


# WINDOW AND TABS
root = Tk()
root.resizable(width=False, height=False)
root.geometry('1000x600')
root.title('Advanced Russian-English-Hebrew translator')
root['bg'] = '#FFE4E1'


tab_control = ttk.Notebook(root)

tab_mini = ttk.Frame(tab_control)
tab_counter = ttk.Frame(tab_control)
tab_translator = ttk.Frame(tab_control)

tab_control.add(tab_mini, text='  MINI TRANSLATOR  ')
tab_control.add(tab_counter, text='  WORD COUNTER  ')
tab_control.add(tab_translator, text='  TRANSLATOR FROM FILE  ')

tab_control.pack(expand=1, fill='both')


# GENERAL FUNCTIONS

def open_file():
    """ loading text from a file for subsequent actions
        can upload files of two formats .doc & .txt
        return str """

    path_to_file = filedialog.askopenfilename()

    if path_to_file.endswith(('.docx', '.doc')):
        doc = docx.Document(path_to_file)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
            text.append('\n\n')
        text = ''.join(text)

    elif path_to_file.endswith('.txt'):
        file = open(path_to_file, 'r', encoding='utf-8')
        text = file.read()
        file.close()

    return text


translator = Translator()

def get_translate(lang, text):
    """ general google - translator """

    words = text.get('1.0', END)
    langer = lang.get()

    if langer == 1:
        a = translator.translate(words, dest='en')
    elif langer == 2:
        a = translator.translate(words, dest='ru')
    else:
        a = translator.translate(words, dest='iw')

    return a


def clear_all():
    """ clear all field in application """

    text_mini_input.delete('1.0', END)
    text_mini_output.delete('1.0', END)
    text_counter_input.delete('1.0', END)
    text_translator_input.delete('1.0', END)
    text_translator_output.delete('1.0', END)
    translator_name_file.delete('1.0', END)


# MINI-TRANSLATOR

def mini_translate(language_mini, text_mini_input):
    """ translate the text_input and paste it into the field text_output """

    a = get_translate(language_mini, text_mini_input)
    text_mini_output.delete('1.0', END)
    text_mini_output.insert('1.0', a.text)


def sound_active(language_mini):
    """ speech synthesizer """

    language = language_mini.get()
    if language == 1:
        word = text_mini_output.get('1.0', END)
    else:
        word = text_mini_input.get('1.0', END)
    engine = pyttsx3.init()
    engine.setProperty('rate', 120)
    engine.setProperty('volume', 1)
    engine.say(word)
    engine.runAndWait()


# radiobuttons for choice of language for translate
language_mini = IntVar(value=1)
language_mini_to_eng = Radiobutton(tab_mini, text="to ENG", value=1, variable=language_mini,
                                   font='Consolas 13', fg='#48494f', padx=15, pady=5)
language_mini_to_eng.place(x=15, y=40)

language_mini_to_ru = Radiobutton(tab_mini, text="to RU", value=2, variable=language_mini,
                                  font='Consolas 13', fg='#48494f', padx=15, pady=5)
language_mini_to_ru.place(x=15, y=70)

language_mini_to_heb = Radiobutton(tab_mini, text="to HEB", value=3, variable=language_mini,
                                   font='Consolas 13', fg='#48494f', padx=15, pady=5)
language_mini_to_heb.place(x=15, y=100)


# input and output fields for text
text_mini_input = Text(tab_mini, font='Consolas 12', width=80, height=6, bg='#ffffff', fg='#3d3d42')
text_mini_input.place(relx=0.5, y=100, anchor=CENTER)

text_mini_output = Text(tab_mini, font='Consolas 12', width=80, height=6, bg='#ffffff', fg='#3d3d42')
text_mini_output.place(relx=0.5, y=260, anchor=CENTER)

# button for translator and synthesizer
btn_mini_translate = Button(tab_mini, width=40, text='TRANSLATE',
                            command=lambda: mini_translate(language_mini, text_mini_input))
btn_mini_translate.place(relx=0.5, y=180, anchor=CENTER)

btn_mini_sound = Button(tab_mini, text='HEAR', width='40', compound=CENTER, command=lambda: sound_active(language_mini))
btn_mini_sound.place(relx=0.5, y=350, anchor=CENTER)


# ADVANCED TRANSLATOR

def get_text():
    text = open_file()
    text_translator_input.delete('1.0', END)
    text_translator_input.insert('1.0', text)


def file_translate(language_translator, text_translator_input):
    a = get_translate(language_translator, text_translator_input)
    text_translator_output.delete('1.0', END)
    text_translator_output.insert('1.0', a.text)


def Export_File():
    dir_name = filedialog.askdirectory()
    os.chdir(dir_name)


def save_file(text_translator_input, text_translator_output):
    text_translator_input = text_translator_input.get('1.0', END)
    text_translator_output = text_translator_output.get('1.0', END)

    doc = docx.Document()
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    row = table.rows[0]
    row.cells[0].text = text_translator_input
    row.cells[1].text = text_translator_output

    file_name = str(translator_name_file.get('1.0', END)).strip() + '.docx'
    if file_name == '.docx':
        msg = "Write the file name in the field next to the button"
        messagebox.showinfo("File saving", msg)
    else:
        Export_File()
        doc.save(file_name)
        msg = f"File {file_name} created successfully"
        messagebox.showinfo("File saving", msg)

# radiobuttons for choice of language for advanced translate
language_translator = IntVar(value=2)
language_translator_to_eng = Radiobutton(tab_translator, text="to ENG", value=1, variable=language_translator, font='Consolas 10', fg='#48494f', padx=15, pady=5)
language_translator_to_eng.place(x=430, y=2)

language_translator_to_ru = Radiobutton(tab_translator, text="to RU", value=2, variable=language_translator, font='Consolas 10', fg='#48494f', padx=15, pady=5)
language_translator_to_ru.place(x=430, y=25)

language_translator_to_heb = Radiobutton(tab_translator, text="to HEB", value=3, variable=language_translator, font='Consolas 10', fg='#48494f', padx=15, pady=5)
language_translator_to_heb.place(x=430, y=50)


# fields for text
text_translator_input = scrolledtext.ScrolledText(tab_translator, font='Consolas 10',
                                   width=65, height=29, bg='#ffffff', fg='#3d3d42', relief='groov', wrap=WORD)
text_translator_input.place(x=15, y=80)

text_translator_output = scrolledtext.ScrolledText(tab_translator, font='Consolas 10',
                                   width=65, height=29, bg='#ffffff', fg='#3d3d42', relief='groov', wrap=WORD)
text_translator_output.place(x=500, y=80)

translator_name_file = Text(tab_translator, font='Consolas 15', width=30, height=1, bg='#ffffff', fg='#3d3d42', relief='groov')
translator_name_file.place(x=320, y=536)


# buttons for load text, translate, save file and clear of fields
btn_translator_load_text = Button(tab_translator, width=40, text='OPEN FILE (docx, txt)', command=get_text)
btn_translator_load_text.place(x=80, y=20)

btn_translator_translate = Button(tab_translator, width=40, text='TRANSLATE',
                                  command=lambda: file_translate(language_translator, text_translator_input))
btn_translator_translate.place(x=600, y=20)

btn_translator_clear = Button(tab_translator, width=15, text='CLEAR ALL', command=clear_all)
btn_translator_clear.place(x=842, y=536)

btn_translator_save_file = Button(tab_translator, width=18, text='SAVE IN NEW FILE:',
                                    command=lambda: save_file(text_translator_input, text_translator_output))
btn_translator_save_file.place(x=170, y=536)


# COUNTER OF WORDS IN TEXT

def loud_text():
    """ loading text from a file in text_counter_input """
    line = open_file()
    text_counter_input.delete('1.0', END)
    text_counter_input.insert('1.0', line)


def get_wordcounter():
    """ clear text from characters,
        calculation of the number of unique words in the text,
        return dict {word: count}"""

    line = text_counter_input.get('1.0', END)
    esc = '0123456789<=>?@[\\]^_`{|}~+-=—)(*&^%$#@!?‘’.,:;"–\'/«»“”'

    for i in esc:
        if i in line:
            line = line.replace(i, '')
    line = line.lower().strip().split()

    wordcounter = dict()

    for word in line:
        if word in wordcounter:
            wordcounter[word] += 1
        else:
            wordcounter[word] = 1

    sort = sorting.get()
    if sort == 1:
        wordcounter = sorted(wordcounter.items(), key=lambda x: x[0])
    else:
        wordcounter = sorted(wordcounter.items(), key=lambda x: x[1], reverse=True)
    wordcounter = dict(wordcounter)

    return wordcounter


def print_wordcounter(sorting):
    """ print a list of unique words based on sorting"""

    wordcounter = get_wordcounter()

    tab_wordcounter = []
    for i in wordcounter:
        tab_wordcounter.append((i, wordcounter[i]))

    columns = ("word", "count")

    tree = ttk.Treeview(tab_counter, height=18, columns=columns, show="headings")
    tree.place(x=550, y=150)

    tree.heading("word", text="WORD")
    tree.heading("count", text="COUNT")

    for i in tab_wordcounter:
        tree.insert("", END, values=i)


def load_into_translater():
    """ loading a list of unique words into the Translator """

    words = get_wordcounter()
    words = "\n".join(list(words))
    text_translator_input.delete('1.0', END)
    text_translator_input.insert('1.0', words)


# field fo load text
text_counter_input = scrolledtext.ScrolledText(tab_counter, font='Consolas 12',
                                   width=50, height=25, bg='#ffffff', fg='#3d3d42', relief='groov', wrap=WORD)
text_counter_input.place(x=50, y=50)


# sort options
sort_label = Label(tab_counter, text="Choosing sorting:", font='Consolas 12', fg='#48494f', padx=15, pady=5)
sort_label.place(x=550, y=10)

sorting = IntVar(value=1)
sort_alphabet = Radiobutton(tab_counter, text="by alphabet", value=1, variable=sorting,
                            font='Consolas 12', fg='#48494f', padx=15, pady=5)
sort_alphabet.place(x=550, y=40)

sort_frequency = Radiobutton(tab_counter, text="by frequency", value=2, variable=sorting,
                             font='Consolas 12', fg='#48494f', padx=15, pady=5)
sort_frequency.place(x=550, y=70)


# buttons for load text, counting, transfer to translator and clear of field
btn_counter_textinput = Button(tab_counter, width=25, text='LOUD TEXT FROM FILE', command=loud_text)
btn_counter_textinput.place(x=50, y=20)

btn_counter_count = Button(tab_counter, width=25, text='COUNT UNIQUE WORDS', command=lambda: print_wordcounter(sorting))
btn_counter_count.place(x=550, y=110)

btn_counter_clear = Button(tab_counter, width=15, text='CLEAR ALL', command=clear_all)
btn_counter_clear.place(x=389, y=20)

btn_counter_to_translator = Button(tab_counter, width=25, text='LOAD INTO TRANSLATOR', command=load_into_translater)
btn_counter_to_translator.place(x=765, y=110)


root.mainloop()

